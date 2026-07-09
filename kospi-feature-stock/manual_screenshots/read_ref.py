import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

d = Document("D:/a2m/atom-harness-g2b/BidAI_사용자매뉴얼_v1.0.docx")

print("=== STYLES ===")
for s in d.styles:
    if s.type is not None:
        try:
            print(f"  {s.name} ({s.type})")
        except Exception:
            pass

print("\n=== SECTIONS ===")
for i, sec in enumerate(d.sections):
    print(f"  section {i}: pw={sec.page_width}, ph={sec.page_height}, lm={sec.left_margin}, rm={sec.right_margin}, tm={sec.top_margin}, bm={sec.bottom_margin}")

print("\n=== HEADINGS AND FIRST 100 PARAGRAPHS ===")
for i, p in enumerate(d.paragraphs[:120]):
    text = p.text.strip()
    style = p.style.name if p.style else ""
    if text:
        print(f"  [{i}] ({style}) {text[:100]}")

print(f"\n=== TOTAL PARAGRAPHS: {len(d.paragraphs)} ===")
print(f"=== TOTAL TABLES: {len(d.tables)} ===")

# sample a heading list
print("\n=== ALL HEADINGS ===")
for p in d.paragraphs:
    if p.style and p.style.name.startswith("Heading"):
        print(f"  ({p.style.name}) {p.text[:120]}")
