# -*- coding: utf-8 -*-
"""Quant Eye 사용자 매뉴얼 v1.0 생성 스크립트."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "D:/a2m/atom-harness-base-Dart/kospi-feature-stock"
SHOT = f"{BASE}/manual_screenshots"
OUT  = f"{BASE}/Quant_Eye_사용자매뉴얼_v1.0.docx"

# ─────────────────────── 문서 초기화 ───────────────────────
doc = Document()

# 페이지: A4, 여백 2.5cm
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# 기본 폰트: 맑은 고딕
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)
rpr = style.element.rPr
if rpr is None:
    rpr = OxmlElement("w:rPr"); style.element.insert(0, rpr)
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts"); rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "맑은 고딕")
rFonts.set(qn("w:ascii"), "맑은 고딕")
rFonts.set(qn("w:hAnsi"), "맑은 고딕")

# Heading 폰트 컬러/사이즈 (한글 대응)
for level, size, color in [(1, 18, RGBColor(0x1F, 0x49, 0x7D)),
                           (2, 14, RGBColor(0x2E, 0x74, 0xB5)),
                           (3, 12, RGBColor(0x1F, 0x49, 0x7D))]:
    s = doc.styles[f"Heading {level}"]
    s.font.name = "맑은 고딕"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    rpr = s.element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr"); s.element.insert(0, rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")


# ─────────────────────── 헬퍼 ───────────────────────
def p(text: str = "", *, bold: bool = False, size: int | None = None,
      align=None, color=None, italic: bool = False) -> None:
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "맑은 고딕"
    r = run._r
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def h(level: int, text: str) -> None:
    doc.add_heading(text, level=level)


def bullet(items: list[str]) -> None:
    for it in items:
        par = doc.add_paragraph(it, style="List Bullet")
        for run in par.runs:
            run.font.name = "맑은 고딕"


def num(items: list[str]) -> None:
    for it in items:
        par = doc.add_paragraph(it, style="List Number")
        for run in par.runs:
            run.font.name = "맑은 고딕"


def mono(text: str) -> None:
    """고정폭 코드 블록 (ASCII 다이어그램)."""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); run._r.insert(0, rPr)
    # 배경음영
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    rPr.append(shd)


def caption(text: str) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r.font.name = "맑은 고딕"


def image(name: str, caption_text: str, width: float = 6.0) -> None:
    path = f"{SHOT}/{name}.png"
    if os.path.exists(path):
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run()
        try:
            run.add_picture(path, width=Inches(width))
        except Exception as e:
            p(f"[이미지 로드 실패: {name} — {e}]", italic=True)
        caption(caption_text)
    else:
        p(f"[이미지 없음: {name}.png]", italic=True)


def table(header: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    tb = doc.add_table(rows=1 + len(rows), cols=len(header))
    tb.style = "Light Grid Accent 1"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tb.rows[0].cells
    for i, txt in enumerate(header):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        run = par.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"
        # 헤더 배경
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows, start=1):
        cells = tb.rows[r_idx].cells
        for i, txt in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            run = par.add_run(str(txt))
            run.font.size = Pt(9.5)
            run.font.name = "맑은 고딕"
    if col_widths:
        for row in tb.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)


def page_break() -> None:
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════
p("")
p("")
p("")
p("Quant Eye", size=42, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
  color=RGBColor(0x1F, 0x49, 0x7D))
p("KOSPI/KOSDAQ 실시간 특징주 탐지 및 매매 추천 시스템",
  size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
  color=RGBColor(0x40, 0x40, 0x40))
p("")
p("")
p("")
p("사용자 매뉴얼", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p("")
p("")
p("")
p("")
p("")
p("")
p("버전: v1.0    |    작성일: 2026-07-09",
  size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
  color=RGBColor(0x60, 0x60, 0x60))
p("Quant Eye — Real-time Feature Stock Detection & AI Recommendation",
  size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
  color=RGBColor(0x60, 0x60, 0x60))
page_break()

# ═══════════════════════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════════════════════
h(1, "목차")
toc_items = [
    "1장. 시스템 개요",
    "   1.1 시스템 소개 및 목적",
    "   1.2 주요 기능 요약",
    "   1.3 시스템 구성도",
    "   1.4 데이터 흐름도",
    "   1.5 용어 정의",
    "",
    "2장. 시스템 설치 및 실행",
    "   2.1 사전 요구사항",
    "   2.2 다른 PC에서 독립 실행하는 방법",
    "   2.3 초기 설정 및 데이터 적재",
    "",
    "3장. 화면별 기능 설명",
    "   3.1 대시보드 (Dashboard)",
    "   3.2 특징주 탐지 (Feature Events)",
    "   3.3 매매 추천 (Recommendations)",
    "   3.4 종목 검색 (Stock Search)",
    "   3.5 성과 추적 (Performance Tracking)",
    "   3.6 모델 성능 (Model Performance)",
    "   3.7 종목 스크리너 (Screener)",
    "   3.8 백테스트 (Backtest)",
    "   3.9 시스템 헬스 (System Health)",
    "   3.10 인텔리전스 (Intel)",
    "   3.11 자동매매 (Trader)",
    "   3.12 관심종목 (Watchlist)",
    "   3.13 알림 (Notifications)",
    "   3.14 랭킹 (Ranking)",
    "   3.15 설정 (Settings)",
    "",
    "4장. 점수 계산 및 추천 알고리즘",
    "   4.1 특징주 탐지 알고리즘",
    "   4.2 매매 추천 점수 계산식 (rec_score)",
    "   4.3 ML 확률 모델",
    "   4.4 시장 국면 판단",
    "   4.5 리스크 스코어 계산",
    "   4.6 성공 확률 캡",
    "",
    "5장. 사용 시나리오",
    "   5.1 일반 투자자 시나리오 (아침 루틴)",
    "   5.2 단타 트레이더 시나리오",
    "   5.3 시스템 관리자 시나리오",
    "   5.4 데이터 관리자 시나리오",
    "",
    "6장. FAQ 및 문제 해결",
    "7장. 부록 (API 엔드포인트, 환경변수 목록)",
]
for it in toc_items:
    par = doc.add_paragraph(it)
    par.paragraph_format.space_after = Pt(2)
page_break()


# ═══════════════════════════════════════════════════════════
# 1장 시스템 개요
# ═══════════════════════════════════════════════════════════
h(1, "1장. 시스템 개요")

h(2, "1.1 시스템 소개 및 목적")
p("Quant Eye는 KOSPI/KOSDAQ 종목의 실시간 시세, 수급, 공시, 뉴스, 재무 데이터를 "
  "종합 수집·분석해 '특징주'를 탐지하고, ML 기반 매매 신호(추천)를 산출·추적하는 "
  "종합 트레이딩 어시스턴트 시스템입니다.")
p("주요 목표는 다음과 같습니다:")
bullet([
    "장중 급등·거래량 폭발·수급 이상·공시 후 반응 등을 초 단위로 포착",
    "9년치(2018~) 유사 사례 검색 + LightGBM 확률 모델로 성공 확률 계산",
    "시장 국면(강세/중립/약세)을 반영한 리스크 조정 매매 추천",
    "KIS(한국투자증권) API를 통한 자동 매매 실행 (모의/실전)",
    "추천 이후 1일/3일/5일/10일 성과를 자동 추적해 모델을 재학습",
])

h(2, "1.2 주요 기능 요약")
table(
    ["카테고리", "기능", "설명"],
    [
        ["시장 관측", "대시보드", "지수·상승/하락 상위·섹터 히트맵·신고가·수급 상위 실시간"],
        ["시장 관측", "인텔리전스", "테마·뉴스·공시·이슈를 통합해 오늘의 시장 스토리 제공"],
        ["시장 관측", "랭킹", "거래대금·급등·이벤트 발생 종목 순위"],
        ["탐지",     "특징주 이벤트", "14종 이벤트 실시간 탐지·시그널 점수화"],
        ["탐지",     "스크리너", "RSI·52W신고가·거래량·수급·ML확률·PER/ROE AND 조합 검색"],
        ["추천",     "매매 추천", "ML 확률+유사사례+리스크로 rec_score(1~100)·success_prob 계산"],
        ["추천",     "관심종목", "지정 종목에 대한 시그널·가격 알림 관리"],
        ["실행",     "자동매매(Trader)", "KIS API를 통한 모의/실전 주문·포지션 관리"],
        ["실행",     "알림", "텔레그램/이메일로 추천·이벤트 발생 시 즉시 전송"],
        ["평가",     "성과 추적", "추천 이후 1d/3d/5d/10d 수익률·손절/익절 히트 추적"],
        ["평가",     "모델 성능", "AUC·정밀도·재현율·이벤트 유형별 승률 대시보드"],
        ["평가",     "백테스트", "과거 데이터로 특징주+추천 규칙 시뮬레이션"],
        ["관리",     "시스템 헬스", "각 마이크로서비스 상태·수집 지연·큐잉 실시간 모니터링"],
        ["관리",     "설정", "임계값·알림·API 키·재학습 옵션 관리"],
    ],
    col_widths=[2.5, 3.5, 10.0],
)

h(2, "1.3 시스템 구성도")
p("Quant Eye는 마이크로서비스 아키텍처를 채택하고, 모든 컨테이너는 docker-compose로 오케스트레이션됩니다.")
mono(
"""┌─────────────────────────────────────────────────────────────────────────────┐
│                        Quant Eye 시스템 구성                                │
├─────────────────────────────────────────────────────────────────────────────┤
│  [외부 데이터]       [수집 레이어]        [처리 레이어]       [서비스 레이어]│
│                                                                             │
│  KIS OpenAPI ─────► collector-tick ──┐                                      │
│  (WebSocket)         (틱/호가)       │                                      │
│                                      ├──► detector ──┐   ┌── api :8000     │
│  KIS REST    ─────► collector-daily ─┤    (14종 탐지) │   │   (FastAPI)     │
│  (일봉/재무)         (일봉/기술지표) │                │   │        ▲        │
│                                      ├──► analyzer ──┤   │        │        │
│  KIND/DART   ─────► collector-news  ─┤    (유사사례)  │   │  [브라우저]     │
│  (공시)                              │                │   │  React SPA      │
│                                      ├──► ml :8001 ──┤   │  (localhost:8000│
│  뉴스 RSS    ─────► collector-supply┤   (LightGBM)   │   │        )        │
│                       (수급 30분)   │                │   │                 │
│                                      ├──► recommender─┴──►│                │
│  금융위 API  ─────► collector-govdata│   (rec_score)      │                │
│  (시총 EOD)         (일 1회 18:00)   │                     │                │
│                                                            │                │
│                     [저장 레이어]                     ├──► trader :8004 ──►│
│                     PostgreSQL 16 + TimescaleDB       │    (KIS 주문 API)  │
│                       + pgvector (유사사례)           │        │           │
│                     Redis 7 (실시간/캐시/Pub-Sub)     │        ▼           │
│                                                       └──► notifier :8003  │
│                     [비동기 채널]                          (Telegram 알림) │
│                     ch:recommendation  ch:feature                          │
└─────────────────────────────────────────────────────────────────────────────┘"""
)

h(2, "1.4 데이터 흐름도")
p("전형적인 신호 발생 흐름:")
mono(
"""1) KIS WebSocket → collector-tick   : 초 단위 호가·체결 저장 (Redis)
2) 30분 배치       → collector-supply : supply_demand 테이블 갱신
3) 이벤트 탐지     → detector          : feature_events 삽입, ch:feature 발행
4) 유사 사례 검색  → analyzer          : pgvector cosine top-K + LightGBM 확률
5) 추천 산출       → recommender       : rec_score + regime 보정 → recommendations
6) 사용자 노출     → api → 프론트엔드  : 대시보드 / 추천 화면
7) 자동 매매(선택)→ trader             : action='BUY' → KIS 주문
8) 성과 추적       → tracking          : 1d/3d/5d/10d 수익률 자동 계산
9) 재학습          → ml-autoretrain    : 14일·28일 00:00 KST 자동 재학습"""
)

h(2, "1.5 용어 정의")
table(
    ["용어", "설명"],
    [
        ["특징주(Feature Event)", "거래량 급증·신고가·수급 이상 등 시장 참여자 관심을 끌 만한 이벤트"],
        ["signal_score", "이벤트별 원시 점수 (0~1). 값이 클수록 강한 신호"],
        ["rec_score", "매매 추천 종합 점수 (1~100). ML+유사사례+리스크 종합"],
        ["success_prob", "추천 성공 확률 (0~0.95). 캡=0.95 (과신 방지)"],
        ["regime", "시장 국면: bull(강세) / neutral(중립) / bear(약세) — KOSPI MA20 기준"],
        ["risk_score", "리스크 점수 (0~1). 값이 클수록 위험 신호(변동성·수급 악화 등)"],
        ["r_1d/3d/5d/10d", "추천 시점 대비 1일/3일/5일/10일 종가 수익률"],
        ["hit_target", "목표가 도달 여부(익절)"],
        ["hit_stop", "손절가 도달 여부"],
        ["dedupe", "종목당 최고 점수 1건만 표시 옵션"],
        ["walk-forward", "시간 순 재학습·검증 방식 (실전 유사)"],
        ["Similar Case", "pgvector로 검색한 과거 유사 이벤트 top-K"],
    ],
    col_widths=[4.0, 12.0],
)
page_break()


# ═══════════════════════════════════════════════════════════
# 2장 시스템 설치 및 실행
# ═══════════════════════════════════════════════════════════
h(1, "2장. 시스템 설치 및 실행")

h(2, "2.1 사전 요구사항")
p("아래 요건을 충족하는 PC에서 원활히 동작합니다.")
table(
    ["항목", "최소 사양", "권장 사양"],
    [
        ["OS",       "Windows 10 (WSL2) / Ubuntu 22.04 / macOS 12", "Windows 11 + WSL2 / Ubuntu 24.04"],
        ["CPU",      "4 코어",    "8 코어 이상"],
        ["메모리",   "8 GB",     "16 GB 이상"],
        ["디스크",   "50 GB SSD","200 GB NVMe SSD"],
        ["Docker",   "Docker Desktop 4.30 이상", "최신 안정 버전"],
        ["Python",   "3.12 (관리 스크립트 사용 시)", "3.12.x"],
        ["Node.js",  "선택 (프론트 재빌드 시)", "20 LTS"],
        ["네트워크","KIS/DART/KIND/금융위 API 접근 가능", "고정 IP + 방화벽 규칙"],
    ],
    col_widths=[2.5, 6.5, 6.5],
)

h(2, "2.2 다른 PC에서 독립 실행하는 방법")
p("빈 PC에 시스템을 새로 설치·실행하는 표준 절차입니다. 순서대로 진행하십시오.")

h(3, "Step 1. Docker Desktop 설치")
num([
    "https://www.docker.com/products/docker-desktop 접속 후 OS에 맞는 설치 파일 다운로드",
    "설치 실행 → WSL2 백엔드 사용 옵션 체크 (Windows)",
    "설치 완료 후 재부팅, Docker Desktop 실행",
    "Settings → Resources → Memory: 7GB 이상 / CPUs: 4 이상 설정",
])

h(3, "Step 2. 프로젝트 소스 복사")
p("다음 중 하나의 방법을 사용합니다.")
bullet([
    "방법 A(Git): git clone <repository-url> kospi-feature-stock",
    "방법 B(ZIP): USB/네트워크 공유로 kospi-feature-stock 폴더 전체 복사",
])
p("복사 후 프로젝트 루트로 이동:")
mono("cd kospi-feature-stock")

h(3, "Step 3. 환경 변수(.env) 파일 설정")
p("프로젝트 루트에 `.env` 파일을 생성하고 아래 항목을 채웁니다.")
mono(
"""# ── PostgreSQL ─────────────────────────────
POSTGRES_DB=fstock
POSTGRES_USER=fstock
POSTGRES_PASSWORD=<임의의 강력한 비밀번호>
POSTGRES_PORT=5432

# ── Redis ─────────────────────────────────
REDIS_PORT=6379

# ── API/Trader 포트 ───────────────────────
API_PORT=8000
ML_API_PORT=8001
NOTIFIER_PORT=8003
TRADER_PORT=8004

# ── KIS OpenAPI (한국투자증권) ────────────
KIS_APP_KEY=<KIS 앱키>
KIS_APP_SECRET=<KIS 앱시크릿>
KIS_ACCOUNT_NO=<계좌번호 8자리>
KIS_ACCOUNT_CD=01
KIS_BASE_URL=https://openapi.koreainvestment.com:9443
KIS_MODE=paper   # paper(모의) / live(실전)

# ── 외부 데이터 ───────────────────────────
DART_API_KEY=<금감원 DART 키>
KIND_API_KEY=<KIND 키>
GOVDATA_API_KEY=<금융위 시가총액 API 키>

# ── 알림 ──────────────────────────────────
TELEGRAM_BOT_TOKEN=
TELEGRAM_CHAT_ID=

# ── ML/추천 임계값 ────────────────────────
REC_COOLDOWN_MINUTES=60
REC_PERF_MIN_PROB=0.236
REGIME_BEAR_THRESHOLD=-3.0
REGIME_BULL_THRESHOLD=-0.5
REGIME_BEAR_SUPPRESS=true"""
)

h(3, "Step 4. Docker Compose로 시스템 기동")
p("아래 명령을 프로젝트 루트에서 실행합니다.")
mono(
"""# 1) 이미지 빌드 + 컨테이너 기동 (백그라운드)
docker compose up -d --build

# 2) 서비스 상태 확인
docker compose ps

# 3) 헬스 체크
curl http://localhost:8000/api/health
curl http://localhost:8001/health   # ML
curl http://localhost:8004/health   # Trader"""
)
p("정상 기동 시 다음 컨테이너가 실행됩니다.")
bullet([
    "fstock-postgres, fstock-redis (인프라)",
    "fstock-collector-tick / -daily / -supply / -news / -batch / -financials / -govdata (수집)",
    "fstock-detector, fstock-analyzer, fstock-ml, fstock-recommender (처리)",
    "fstock-api, fstock-trader, fstock-notifier (서비스)",
])

h(3, "Step 5. 초기 데이터 부트스트랩 (최초 1회)")
p("최초 실행 시 종목 마스터·과거 일봉·재무 데이터를 적재해야 합니다.")
mono(
"""# 종목 마스터 + 최근 일봉 + 재무 초기 로드
python setup_bootstrap.py

# (선택) 9년치 일봉 백필 — 별도 프로파일
docker compose --profile backfill up collector-bars-backfill

# (선택) ML 모델 학습 (walk-forward)
docker compose --profile tools run --rm ml-train"""
)

h(3, "Step 6. 브라우저 접속")
p("정상 기동 후 브라우저에서 아래 주소로 접속합니다.")
bullet([
    "동일 PC: http://localhost:8000",
    "동일 네트워크 다른 PC: http://<서버IP>:8000 (예: http://192.168.1.100:8000)",
])
p("사이드바에 15개 메뉴가 표시되면 정상 설치 완료입니다.")

h(3, "Step 7. 시스템 종료/재시작")
mono(
"""# 정지 (컨테이너는 유지)
docker compose stop

# 완전 종료 (컨테이너 제거)
docker compose down

# 완전 종료 + 데이터 삭제 (주의: DB 데이터 소실)
docker compose down -v

# 재기동
docker compose up -d"""
)

h(2, "2.3 초기 설정 및 데이터 적재")
p("설치 직후에는 데이터가 비어 있어 대시보드/추천이 표시되지 않을 수 있습니다. 아래 순서로 데이터를 채웁니다.")
num([
    "종목 마스터 등록: scripts/init_stocks.py 실행 (KOSPI/KOSDAQ 전 종목)",
    "일봉 백필: docker compose --profile backfill up collector-bars-backfill (BARS_BACKFILL_DAYS 조정)",
    "수급 백필: scripts/backfill_supply.py --days 207",
    "재무 수집: docker compose exec collector-financials python entrypoints/financials_worker.py --once",
    "시가총액 백필: docker compose exec collector-govdata python entrypoints/govdata_worker.py --once",
    "ML 학습: docker compose --profile tools run --rm ml-train",
    "탐지·추천 로직 워밍업: 장 시작 시각까지 약 10~20분 대기 (Redis 캐시가 채워짐)",
])
page_break()


# ═══════════════════════════════════════════════════════════
# 3장 화면별 기능 설명
# ═══════════════════════════════════════════════════════════
h(1, "3장. 화면별 기능 설명")

# 3.1 대시보드
h(2, "3.1 대시보드 (Dashboard)")
image("01_dashboard", "그림 3-1. 대시보드 화면")
p("대시보드는 시장 전체 현황과 오늘의 핵심 신호를 한 페이지에 요약합니다.")
p("표시 항목:")
bullet([
    "KOSPI/KOSDAQ 실시간 지수 (KIS API → 30초 캐시, 실패 시 daily_bars 폴백)",
    "시장 국면 배지: bull(강세) / neutral(중립) / bear(약세) — KOSPI MA20 대비 %",
    "상승/하락 상위 5개 종목 (실시간 호가 보정)",
    "섹터 히트맵: 섹터별 평균 등락률 + 상위 종목 top-5",
    "신고가 종목: 최근 24시간 내 52W/26W/20D 신고가 이벤트",
    "외국인/기관 순매수 상위 10개 (supply_demand + daily_bars 병합)",
])
p("사용 팁:")
bullet([
    "지수 배지가 bear일 때는 추천 확률이 자동으로 0.75배로 감쇄됩니다 (2.4 시장 국면 참조).",
    "섹터 히트맵의 진한 초록/빨강 셀은 클릭 시 해당 섹터의 종목 목록으로 필터링됩니다.",
])
page_break()

# 3.2 특징주 탐지
h(2, "3.2 특징주 탐지 (Feature Events)")
image("02_features", "그림 3-2. 특징주 탐지 화면")
p("실시간으로 탐지된 특징주 이벤트를 시그널 점수(signal_score) 순으로 조회합니다.")
p("주요 필터:")
bullet([
    "이벤트 타입: 14종 중 선택",
    "종목 코드 / 시장(KOSPI/KOSDAQ)",
    "최소 시그널 점수 (기본 0.5)",
    "조회 기간 (시간 단위, 최대 168)",
    "dedupe: 종목당 최고 점수 1건만 표시",
])
p("이벤트 타입별 설명 (자세한 로직은 4.1절 참조):")
table(
    ["이벤트 타입", "설명", "핵심 조건"],
    [
        ["VOLUME_SURGE",         "거래량 급증", "당일 거래량 ÷ 20일 평균 ≥ 3"],
        ["AMOUNT_SURGE",         "거래대금 급증", "당일 거래대금 ÷ 20일 평균 ≥ 3"],
        ["BREAKOUT_52W",         "52주 신고가",  "당일 고가 ≥ 최근 252 거래일 최고가"],
        ["BREAKOUT_26W",         "26주 신고가",  "당일 고가 ≥ 최근 126 거래일 최고가"],
        ["BREAKOUT_13W",         "13주 신고가",  "당일 고가 ≥ 최근 63 거래일 최고가"],
        ["BREAKOUT_20D",         "20일 신고가",  "당일 고가 ≥ 최근 20 거래일 최고가"],
        ["VI_TRIGGERED",         "변동성완화장치 발동", "10% 이상 등락 후 2분 정지"],
        ["LONG_WHITE_CANDLE",    "장대양봉",     "몸통 > 5%, 종가 근접 고가"],
        ["HAMMER_CANDLE",        "망치형",       "아래꼬리 > 몸통×2"],
        ["MORNING_STAR",         "샛별형",       "3봉 반전 패턴"],
        ["SUPPLY_ANOMALY",       "수급 이상",    "외국인+기관 대량 순매수"],
        ["POST_DISCLOSURE_SURGE","공시 후 급등", "공시 시각 이후 30분 내 5%↑"],
        ["SHORT_SURGE",          "공매도 급증",  "공매도 잔고 급등"],
        ["DUAL_BUY_STREAK",      "외인+기관 연속 순매수", "N일 연속 동반 매수"],
    ],
    col_widths=[3.5, 3.5, 9.0],
)
p("이벤트 행 클릭 시 팝업:")
bullet([
    "이벤트 상세 (탐지 시각·가격·거래량·시그널 점수)",
    "유사 사례 5건 (pgvector cosine 검색) + 각 사례의 이후 15일 가격 차트",
    "관련 추천이 있으면 추천 상세로 연결",
])
page_break()

# 3.3 추천
h(2, "3.3 매매 추천 (Recommendations)")
image("03_recommendations", "그림 3-3. 매매 추천 화면")
p("탐지된 특징주 이벤트를 recommender 서비스가 ML 확률과 유사사례로 재평가해 산출한 매매 추천 목록입니다.")
p("주요 필터:")
bullet([
    "액션: ALL / BUY / SELL / SKIP",
    "시장 / 종목 코드",
    "최소 성공 확률(min_prob, 기본 0.30)",
    "조회 기간 (시간 단위, 최대 168)",
    "dedupe: 종목당 최고 확률 1건만 표시 (기본 ON)",
])
p("각 추천 카드에서 확인할 수 있는 정보:")
bullet([
    "종목 코드/이름/시장, 액션(BUY/SELL/SKIP)",
    "성공 확률 success_prob (0.95 캡 적용)",
    "예상 수익률 expected_return, 예상 보유일 expected_hold_days",
    "진입가·목표가·손절가, 리스크/보상 비율 risk_reward_ratio",
    "리스크 점수 risk_score, 총 점수 rec_score(1~100)",
    "판단 근거 rationale (ML 확률·패턴 매칭·수급·시장국면)",
    "유사 사례 리스트 (과거 성공/실패 사례)",
])
p("추천 상태 배지:")
bullet([
    "A등급 (rec_score ≥ 80): 높은 신뢰도",
    "B등급 (60~79): 보통",
    "C등급 (40~59): 낮음 — 유사사례 부족 경고",
    "D등급 (<40): 매우 낮음 — 참고용",
])
p("시장 국면 자동 감쇄:")
bullet([
    "bear (KOSPI가 MA20 대비 -3% 이하): success_prob ×0.75",
    "neutral (-3% ~ -0.5%): success_prob ×0.88",
    "bull (-0.5% 이상): 감쇄 없음",
    "REGIME_BEAR_SUPPRESS=true 시 bear 구간은 추천 자체가 억제됨",
])
page_break()

# 3.4 종목 검색
h(2, "3.4 종목 검색 (Stock Search)")
image("04_stock_search", "그림 3-4. 종목 검색 화면")
p("종목명 또는 코드로 검색하여 해당 종목의 상세 정보를 조회합니다.")
p("표시 정보:")
bullet([
    "기본 정보: 코드·이름·시장·업종·상장주식수·시가총액",
    "일봉 차트 (최근 6개월 캔들 + 20일/60일 이동평균)",
    "실시간 호가 (KIS WebSocket 캐시)",
    "재무 요약: PER·PBR·ROE·EPS (최근 분기 기준)",
    "이벤트 이력: 해당 종목에서 발생한 최근 특징주 이벤트 목록",
    "추천 이력: 해당 종목의 매매 추천 이력과 성과 (1d/3d/5d/10d)",
    "관심종목 추가/삭제 버튼",
])
page_break()

# 3.5 성과 추적
h(2, "3.5 성과 추적 (Performance Tracking)")
image("05_performance_tracking", "그림 3-5. 성과 추적 화면")
p("모든 매매 추천의 성과를 추적하고 승률·수익률을 집계합니다.")
p("탭 구성:")
bullet([
    "활성 추적 (Active): tracking_complete=FALSE인 추천 — 진행 중",
    "완료 이력 (History): r_5d NOT NULL인 완료된 추천",
    "요약 통계 (Summary): 승률·평균수익률·목표가 도달률·손절률",
    "이벤트별 성과 (By Event): 이벤트 유형별 승률·평균 수익률",
])
p("각 행 표시 항목:")
bullet([
    "추천 시점, 종목, 액션, 진입가·목표가·손절가, 성공확률",
    "r_1d / r_3d / r_5d / r_10d: 시점 대비 각 기간 종가 수익률(%)",
    "max_return: 진입 후 최고 수익률",
    "hit_target: 목표가 도달 여부 (익절)",
    "hit_stop: 손절가 도달 여부",
    "is_success: 성공 여부 (기본: r_5d > 0)",
])
p("성과는 매일 장 마감 후 collector-batch가 자동 갱신합니다. days 파라미터는 요약 통계 집계 범위입니다.")
page_break()

# 3.6 모델 성능
h(2, "3.6 모델 성능 (Model Performance)")
image("06_model_performance", "그림 3-6. 모델 성능 화면")
p("ML(LightGBM) 모델의 예측 성능을 시각화합니다.")
p("표시 지표:")
bullet([
    "AUC (Area Under ROC): 최근 학습 시 검증 AUC (예: 0.6091)",
    "정밀도 (Precision) / 재현율 (Recall) / F1 스코어",
    "예측 확률 캘리브레이션 곡선 (predicted vs actual)",
    "피처 중요도 상위 20개 (LightGBM feature_importance)",
    "학습 이력: 재학습 시각·데이터 기간·AUC 추이",
    "이벤트 유형별 정밀도 (14종)",
])
p("최근 학습 정보는 model_cache 볼륨의 metadata.json에서 로드됩니다. 재학습은 매 14일·28일 00:00 KST에 auto_retrain.py가 자동 실행합니다.")
page_break()

# 3.7 스크리너
h(2, "3.7 종목 스크리너 (Screener)")
image("07_screener", "그림 3-7. 종목 스크리너 화면")
p("다중 조건 AND 조합으로 종목을 필터링합니다. 결과는 2분 캐시됩니다.")
p("필터 조건:")
table(
    ["필터", "타입", "설명"],
    [
        ["RSI 최솟값/최댓값",      "0~100",  "Wilder's RSI 14일 기준"],
        ["52W 신고가 N% 이내",     "0~100",  "종가가 52주 최고가의 N% 이내"],
        ["거래량 비율 최솟값",     "float",  "당일 거래량 ÷ 20일 평균"],
        ["외국인 N일 연속 순매수", "1~20",   "supply_demand.foreign_net > 0"],
        ["ML 확률 최솟값",         "0~1",    "signal_data.ml_prob"],
        ["이벤트 타입",            "다중선택","최근 30일 내 발생"],
        ["시장",                   "ALL/KOSPI/KOSDAQ", ""],
        ["PER 상한 / ROE 하한",    "float",  "재무 기반 필터"],
        ["결과 개수",              "1~500",  "기본 50"],
    ],
    col_widths=[4.5, 3.0, 9.0],
)
p("각 결과 행은 매칭된 조건 목록(match_conditions)을 함께 표시하여 왜 필터에 걸렸는지 이해할 수 있습니다.")
page_break()

# 3.8 백테스트
h(2, "3.8 백테스트 (Backtest)")
image("08_backtest", "그림 3-8. 백테스트 화면")
p("과거 데이터로 특징주 탐지 + 추천 규칙을 시뮬레이션하여 전략을 검증합니다.")
p("입력 파라미터:")
bullet([
    "기간(시작·종료 날짜)",
    "이벤트 유형 (선택 다중)",
    "최소 시그널 점수 / 최소 ML 확률",
    "포지션 사이즈(원)·최대 동시 포지션 수",
    "익절/손절 정책 (%p 또는 rec 값 사용)",
    "수수료·세금 (bp) — 매수·매도 각 4bp / 매도 세금 20bp 기본",
])
p("출력:")
bullet([
    "누적 손익 그래프, 자본 곡선",
    "총 거래·승률·평균 수익률·최대 낙폭(MDD)·샤프 지수",
    "이벤트 유형별 성과 breakdown",
    "월별/일별 손익 히트맵",
])
p("성능: pandas 핫루프 제거로 최적화되어 3년치 백테스트가 약 1~2초에 완료됩니다.")
page_break()

# 3.9 시스템 헬스
h(2, "3.9 시스템 헬스 (System Health)")
image("09_system_health", "그림 3-9. 시스템 헬스 화면")
p("각 마이크로서비스의 상태·최근 활동·큐 지연을 실시간으로 확인합니다.")
p("표시 항목:")
bullet([
    "컨테이너 상태 (healthy/unhealthy/exited)",
    "각 서비스의 마지막 이벤트 시각 (heartbeat)",
    "Redis Pub/Sub 채널별 최근 메시지 수",
    "DB 커넥션 풀 사용률",
    "수집 지연: 마지막 틱/일봉/뉴스 저장 시각",
    "오류 카운트 (최근 1시간)",
    "메모리·CPU 사용률 (Prometheus 활성 시)",
])
p("이상 감지 시 화면 상단에 경고 배너가 표시됩니다.")
page_break()

# 3.10 인텔리전스
h(2, "3.10 인텔리전스 (Intel)")
image("10_intel", "그림 3-10. 인텔리전스 화면")
p("오늘의 시장 스토리를 테마·뉴스·공시로 통합해 보여줍니다.")
p("섹션:")
bullet([
    "테마 순위: 최근 상승률·거래대금 기준 상위 테마 + 대표 종목",
    "역전 테마: 최근 5일 대비 오늘 순위 급상승 테마",
    "핵심 공시: 유가증권·코스닥 정정공시·주요 사업보고서",
    "뉴스 하이라이트: 언급 종목 클러스터링 + 감성 분석",
    "이슈 발생 종목: 뉴스+공시+이벤트 3중 겹침 종목",
])
page_break()

# 3.11 자동매매
h(2, "3.11 자동매매 (Trader)")
image("11_trader", "그림 3-11. 자동매매 화면")
p("KIS OpenAPI를 통한 자동 주문 실행 및 포지션 관리 화면입니다.")
p("탭 구성:")
bullet([
    "포지션(Position): 현재 보유 종목·평가 손익·목표/손절 도달 여부",
    "주문 이력(Orders): 오늘 발주된 주문·체결 상태",
    "설정: KIS 모드(paper/live)·최대 동시 포지션·1건당 금액 한도",
    "리스크 가드: 일 손실 한도·연속 손실 후 자동 정지",
])
p("동작 방식:")
num([
    "recommender가 ch:recommendation에 BUY 시그널 발행",
    "trader가 신호 수신 → 리스크 가드 통과 확인",
    "KIS 주문 API 호출 (모의/실전)",
    "체결 결과를 positions 테이블에 저장",
    "매 30초 목표/손절 도달 여부 확인 → 익절/손절 주문",
    "성과를 recommendation_performance에 반영",
])
p("주의: KIS_MODE=live 설정 시 실계좌로 실제 주문이 발주됩니다. 반드시 paper 모드로 충분히 검증 후 전환하십시오.")
page_break()

# 3.12 관심종목
h(2, "3.12 관심종목 (Watchlist)")
image("12_watchlist", "그림 3-12. 관심종목 화면")
p("사용자가 지정한 종목을 별도 리스트로 관리하고 알림을 설정합니다.")
p("기능:")
bullet([
    "관심 종목 추가/삭제 (다중 폴더 지원)",
    "종목별 알림 조건: 가격 도달·이벤트 발생·추천 발생",
    "일괄 조회: 현재가·등락률·오늘의 이벤트·최근 추천",
    "메모 필드: 매매 아이디어·근거 기록",
])
page_break()

# 3.13 알림
h(2, "3.13 알림 (Notifications)")
image("13_notifications", "그림 3-13. 알림 화면")
p("시스템이 발송한 모든 알림 이력을 조회합니다.")
p("표시 정보:")
bullet([
    "알림 시각·채널(telegram/email)·수신자·상태(전송/실패)",
    "알림 유형: 매매 추천·특징주 이벤트·시스템 경고·리스크 가드 발동",
    "내용 미리보기 + 원본 페이로드",
    "재전송 버튼 (실패 알림 대상)",
])
p("알림 채널 설정은 3.15 설정 화면에서 관리합니다.")
page_break()

# 3.14 랭킹
h(2, "3.14 랭킹 (Ranking)")
image("14_ranking", "그림 3-14. 랭킹 화면")
p("다양한 지표별 종목 순위를 제공합니다.")
p("랭킹 유형:")
bullet([
    "거래대금 상위 / 거래량 상위",
    "상승률 상위 / 하락률 상위",
    "이벤트 발생 건수 상위 (최근 24시간)",
    "외국인 순매수 상위 / 기관 순매수 상위",
    "ML 확률 상위",
    "rec_score 상위",
])
p("각 랭킹은 최대 100위까지 표시되며, 시장별(KOSPI/KOSDAQ) 필터가 가능합니다.")
page_break()

# 3.15 설정
h(2, "3.15 설정 (Settings)")
image("15_settings", "그림 3-15. 설정 화면")
p("시스템 파라미터를 관리자 화면에서 조정합니다.")
p("설정 카테고리:")
bullet([
    "탐지 임계값: 이벤트별 시그널 점수 최소치·감쇄 계수",
    "추천 옵션: min_prob·쿨다운(분)·시장국면 감쇄 배율",
    "알림: 텔레그램 봇 토큰/채팅 ID·이메일 SMTP·수신 조건",
    "자동매매: KIS 앱키·계좌·모드·일 손실 한도",
    "재학습: 자동 재학습 요일(14일·28일)·AUC 임계 (0.57)",
    "데이터 관리: 백필 트리거·재수집·오래된 데이터 정리",
])
p("변경한 설정은 즉시 Redis에 반영되어 재기동 없이 적용됩니다. 단, 이미지 재빌드가 필요한 항목(포트 등)은 docker compose up -d --build 재실행이 필요합니다.")
page_break()


# ═══════════════════════════════════════════════════════════
# 4장 점수 계산 및 추천 알고리즘
# ═══════════════════════════════════════════════════════════
h(1, "4장. 점수 계산 및 추천 알고리즘")

h(2, "4.1 특징주 탐지 알고리즘 (이벤트 타입별)")
p("각 이벤트는 detector 서비스의 rules/ 모듈에서 독립적으로 평가됩니다. 기본 조건 예시:")
mono(
"""VOLUME_SURGE       : today.volume / avg20 >= VOLUME_SURGE_RATIO(default 3.0)
                    AND today.volume >= VOLUME_MIN_ABS(default 100,000)
AMOUNT_SURGE       : today.amount / avg20_amount >= AMOUNT_SURGE_RATIO(default 3.0)
BREAKOUT_52W/26W/13W/20D
                    : today.high >= MAX(high[len_days])
LONG_WHITE_CANDLE  : (close - open)/open >= 0.05 AND (high - close)/close < 0.01
HAMMER_CANDLE      : lower_shadow > 2 * body_size AND upper_shadow < body_size
MORNING_STAR       : 3봉 반전 (하락봉 → 도지 → 강한 상승봉)
VI_TRIGGERED       : KIS VI 알림 수신 시 즉시
SUPPLY_ANOMALY     : abs(foreign_net + inst_net) >= SUPPLY_STD_MULT(3) * σ
POST_DISCLOSURE_SURGE : 공시 후 30분 이내 change_rate >= 5%
SHORT_SURGE        : short_balance / avg30 >= 2.0
DUAL_BUY_STREAK    : N일 연속 foreign_net>0 AND inst_net>0"""
)
p("각 이벤트의 signal_score(0~1)는 다음 요소로 계산됩니다:")
bullet([
    "기본 조건 초과 비율 (예: 거래량 3배 초과 → 0.5, 5배 → 0.75, 10배 → 1.0)",
    "가격 반응 (change_rate, 절대값 클수록 점수↑)",
    "수급 강도 (외국인/기관 순매수 규모)",
    "규모 보정 (log(market_cap) 기반 가중)",
])

h(2, "4.2 매매 추천 점수 계산식 (rec_score)")
p("recommender 서비스의 entry_recommender._compute_rec_score() 로직:")
mono(
"""rec_score = ml_component + pattern_component + return_adj − risk_penalty
            (최종 정수, 1~100 사이로 clip)

① ml_component      = (ml_prob / 0.95) × 55                     # 최대 55점
② pattern_component = (min(sim_prob_raw, 0.93) / 0.93) × 30 × w  # 최대 30점
                     w = min(1.0, n_cases / 30.0)   # 유사사례 30건 이상 = 풀 가중
③ return_adj                                                     # -20 ~ +15
     avg_sim_return ≥ 5.0%   → +15
     avg_sim_return ≥ 0.0%   → +8
     avg_sim_return ≥ -3.0%  → 0
     avg_sim_return ≥ -7.0%  → -10
     그 외 (매우 나쁨)       → -20
④ risk_penalty      = min(30, len(risk_factors) × 10)            # 최대 -30점"""
)
p("등급 매핑:")
bullet([
    "A (≥80): 높음 — 3요소 모두 강력",
    "B (60~79): 보통 — 대부분의 매수 신호",
    "C (40~59): 낮음 — 유사사례 부족 경고",
    "D (<40): 매우 낮음 — 참고용",
])

h(2, "4.3 ML 확률 모델")
p("Quant Eye의 매매 성공 확률(success_prob)은 LightGBM 이진 분류 모델로 계산됩니다.")
p("모델 사양:")
bullet([
    "알고리즘: LightGBM (gradient boosting)",
    "학습: walk-forward (train 2020~2023 / val 2024 / test 2025)",
    "레이블: relative — 이벤트 발생 후 5일 수익률이 KOSPI 대비 +2%p 이상",
    "피처 수: 71개 (2026-07 기준)",
    "SMOTE로 클래스 불균형 완화",
    "Optuna 30 trial 하이퍼파라미터 튜닝",
    "최근 검증 AUC: 0.6091",
])
p("주요 피처 카테고리:")
mono(
"""가격/거래량: close, open, high, low, volume, amount, change_rate,
              volume_ratio_20, amount_ratio_20, gap_pct
기술 지표:    rsi_14, ma5, ma20, ma60, ma_ratio, macd, bollinger_pos, atr
수급:         foreign_net_5d, foreign_net_20d, inst_net_5d, inst_net_20d,
              retail_net_5d, foreign_hold_rate, short_ratio
재무:         per, pbr, roe, eps_growth, log_market_cap, sector_onehot
시장 국면:    kospi_ma20_dist, vix_kr, regime_dummies
공시/뉴스:    days_since_disclosure, news_sentiment, news_count_7d
과거 이벤트:  event_count_30d, prior_success_rate, similar_case_avg_ret"""
)
p("추론 시 ml_prob는 캘리브레이션 후 recommender로 전달되어 최종 success_prob로 저장됩니다.")

h(2, "4.4 시장 국면 판단 (KOSPI MA20 기준)")
p("KOSPI 지수의 20일 이동평균 대비 현재가 위치로 3단계 국면을 판정합니다.")
mono(
"""pct_from_ma20 = (kospi_close − ma20) / ma20 × 100

phase = 'bear'    if pct_from_ma20 <  REGIME_BEAR_THRESHOLD (-3.0)
      = 'neutral' if pct_from_ma20 <  REGIME_BULL_THRESHOLD (-0.5)
      = 'bull'    otherwise

# 국면별 성공 확률 감쇄
success_prob *= 1.00 (bull)
             *= 0.88 (neutral)  # REGIME_NEUTRAL_PROB_MULT
             *= 0.75 (bear)     # REGIME_BEAR_PROB_MULT

# REGIME_BEAR_SUPPRESS=true 이면 bear 국면에서는 추천 발행 자체를 건너뜀"""
)
p("regime 정보는 30분 단위로 Redis(market:regime)에 캐시되고 대시보드/추천 화면에 배지로 표시됩니다.")

h(2, "4.5 리스크 스코어 계산")
p("risk_score(0~1)는 다음 요소를 합산한 뒤 정규화됩니다.")
bullet([
    "변동성(ATR 20일) / 가격 × 100 → 값이 클수록 위험",
    "외국인/기관 순매도 강도",
    "공매도 잔고 급증",
    "최근 5일 최대 낙폭",
    "테마 과열 지표 (연관 종목 동시 급등 후 조정)",
    "재무 취약 (PBR>10 or 부채비율>200%)",
])
p("risk_reward_ratio = (target - entry) / (entry - stop_loss) 로 계산되며 2.0 미만이면 SKIP 처리됩니다.")

h(2, "4.6 성공 확률 캡 (0.95)")
p("과신 방지를 위해 모든 success_prob 값은 하한 0.0, 상한 0.95로 clip됩니다.")
mono(
"""_MAX_PROB = 0.95   # recommendation_service.py
success_prob = min(_MAX_PROB, float(raw_prob))"""
)
p("이는 극단적으로 높은 예측치가 실제로는 캘리브레이션되지 않는 문제를 완화하기 위한 안전장치입니다.")
page_break()


# ═══════════════════════════════════════════════════════════
# 5장 사용 시나리오
# ═══════════════════════════════════════════════════════════
h(1, "5장. 사용 시나리오")

h(2, "5.1 일반 투자자 시나리오 (아침 루틴)")
p("장 개시 전(오전 8:30~9:00)에 시장 상황을 파악하고 오늘 관심 종목을 선정하는 흐름입니다.")
num([
    "브라우저에서 http://localhost:8000 접속 → 대시보드 확인",
    "지수 배지(bull/neutral/bear) 및 섹터 히트맵으로 오늘 강세/약세 섹터 파악",
    "'신고가 종목' 카드에서 어제 종가 기준 신고가 이벤트 확인",
    "'매매 추천' 메뉴 → dedupe ON, min_prob=0.40, action=BUY 필터",
    "rec_score 80점 이상(A등급) 종목 3~5개를 관심종목으로 등록",
    "장 시작 후 관심종목 화면에서 실시간 등락 관찰",
    "익절/손절 도달 시 알림(텔레그램) 수신 → 판단",
])

h(2, "5.2 단타 트레이더 시나리오")
p("장중 실시간 이벤트를 포착해 짧은 호흡의 매매를 진행합니다.")
num([
    "'특징주 탐지' 화면 → hours=1, min_score=0.7로 최근 1시간 강한 신호만 조회",
    "VOLUME_SURGE + BREAKOUT_20D가 동시에 발생한 종목 우선 확인",
    "이벤트 클릭 → 유사사례 5건의 이후 15일 차트 검토",
    "매매 추천 화면에서 해당 종목의 추천 상세 확인",
    "risk_reward_ratio ≥ 2.5, success_prob ≥ 0.55 이면 진입 고려",
    "자동매매를 사용 중이라면 trader가 자동으로 주문 발주",
    "성과 추적 화면에서 실시간 손익 확인, 목표/손절 도달 시 자동 청산",
])

h(2, "5.3 시스템 관리자 시나리오")
p("주간 시스템 점검·이슈 대응 흐름입니다.")
num([
    "매일 오전 시스템 헬스 대시보드 확인 → 모든 서비스 healthy 상태 검증",
    "수집 지연 5분 초과 시 해당 collector 재시작 (docker compose restart collector-xxx)",
    "매주 월요일 모델 성능 화면에서 AUC 하락(<0.57) 여부 점검",
    "AUC 하락 시 수동 재학습 트리거 (docker compose --profile tools run --rm ml-train)",
    "매월 백테스트 화면에서 최근 30일 파라미터 검증",
    "설정 화면에서 임계값 미세 조정 후 A/B 비교",
])

h(2, "5.4 데이터 관리자 시나리오")
p("데이터 정합성·백필·재수집 흐름입니다.")
num([
    "시스템 헬스에서 특정 종목 last_date가 오래된 경우 확인",
    "docker compose exec collector-daily python entrypoints/daily_bar_worker.py --code 005930 --once 로 개별 재수집",
    "전체 백필: docker compose --profile backfill up collector-bars-backfill (BARS_BACKFILL_DAYS 조정)",
    "재무 재수집: docker compose exec collector-financials python entrypoints/financials_worker.py --once",
    "시가총액 재백필: collector-govdata 워커 수동 실행",
    "DB 진단: docker compose exec postgres psql -U fstock -d fstock -c 'SELECT COUNT(*) FROM daily_bars;'",
    "인덱스/HNSW 재구축: infra/postgres/V*.sql 스크립트 수동 적용",
])
page_break()


# ═══════════════════════════════════════════════════════════
# 6장 FAQ 및 문제 해결
# ═══════════════════════════════════════════════════════════
h(1, "6장. FAQ 및 문제 해결")

table(
    ["증상", "원인", "해결"],
    [
        ["대시보드에 데이터가 표시되지 않음",
         "초기 데이터 미적재 또는 collector 미기동",
         "python setup_bootstrap.py 실행 후 docker compose ps로 collector 상태 확인"],
        ["/api/health가 500 오류",
         "postgres/redis가 아직 준비 중",
         "docker compose logs postgres | tail 로 확인, 30초~1분 대기 후 재시도"],
        ["추천이 하나도 뜨지 않음",
         "bear 국면에서 REGIME_BEAR_SUPPRESS=true",
         ".env에서 false로 변경 후 recommender 재시작, 또는 시장 국면 대기"],
        ["ML 확률이 0.5에 몰림",
         "모델 미학습 또는 피처 결측",
         "docker compose --profile tools run --rm ml-train 실행 후 model_cache 확인"],
        ["자동매매가 주문을 내지 않음",
         "KIS 토큰 만료 또는 리스크 가드 발동",
         "trader 로그 확인, 필요시 KIS 앱키 재발급 → .env 갱신"],
        ["섹터 히트맵이 비어 있음",
         "stocks.sector 미기입",
         "scripts/refresh_sector.py 실행"],
        ["프론트엔드 빌드가 자산과 불일치",
         "assets 파일 삭제/이동",
         "cd frontend && npm run build 재실행, api 이미지 재빌드"],
        ["백테스트가 매우 느림",
         "pandas 핫루프 잔존",
         "최신 코드(pull) 확인, 커밋 dff88bd 이후 84× 최적화 반영됨"],
        ["Redis 메모리 부족",
         "실시간 틱 캐시 폭증",
         "redis.conf의 maxmemory 상향 또는 collector-tick의 TTL 조정"],
        ["Postgres OOM",
         "shared_buffers 부족",
         "docker-compose.yml의 postgres 리소스 limits 상향"],
    ],
    col_widths=[5.0, 5.0, 6.5],
)
page_break()


# ═══════════════════════════════════════════════════════════
# 7장 부록
# ═══════════════════════════════════════════════════════════
h(1, "7장. 부록")

h(2, "부록 A. 주요 API 엔드포인트")
table(
    ["엔드포인트", "메서드", "설명"],
    [
        ["/api/health",                              "GET",  "API 헬스체크"],
        ["/api/market/summary",                      "GET",  "KOSPI/KOSDAQ 등락 요약"],
        ["/api/market/movers?limit=10",              "GET",  "상승/하락 상위"],
        ["/api/market/foreign-flow",                 "GET",  "외국인/기관 순매수 상위"],
        ["/api/market/index-live",                   "GET",  "실시간 지수 (KIS)"],
        ["/api/market/sector-heatmap",               "GET",  "섹터 히트맵"],
        ["/api/market/new-highs?hours=24",           "GET",  "신고가 종목"],
        ["/api/market/overview",                     "GET",  "대시보드 종합 응답"],
        ["/api/features?event_type=&hours=72",       "GET",  "특징주 이벤트 목록"],
        ["/api/features/types",                      "GET",  "지원 이벤트 타입 14종"],
        ["/api/features/today/summary",              "GET",  "오늘의 이벤트 요약"],
        ["/api/features/{id}/similar?top_k=10",      "GET",  "유사 사례"],
        ["/api/features/{id}/similar-with-bars",     "GET",  "유사 사례 + 캔들"],
        ["/api/recommendations?dedupe=true",         "GET",  "매매 추천 목록"],
        ["/api/recommendations/buy",                 "GET",  "BUY 추천 최근 20건"],
        ["/api/recommendations/stats/performance",   "GET",  "성과 통계"],
        ["/api/recommendations/performance/summary", "GET",  "성과 요약"],
        ["/api/recommendations/performance/history", "GET",  "완료 이력"],
        ["/api/recommendations/performance/by-event","GET",  "이벤트별 승률"],
        ["/api/recommendations/{code}/signals",      "GET",  "종목별 신호 이력"],
        ["/api/recommendations/{code}/latest",       "GET",  "종목별 최신 추천"],
        ["/api/screener/run",                        "POST", "스크리너 실행"],
        ["/api/stocks/search?q=삼성",                "GET",  "종목 검색"],
        ["/api/stocks/{code}",                       "GET",  "종목 상세"],
        ["/api/tracking/*",                          "GET",  "성과 추적"],
        ["/api/ml/status",                           "GET",  "모델 상태"],
        ["/api/ml/features/importance",              "GET",  "피처 중요도"],
        ["/api/backtest/run",                        "POST", "백테스트 실행"],
        ["/api/watchlist",                           "*",    "관심종목 CRUD"],
        ["/api/notifications",                       "GET",  "알림 이력"],
        ["/api/ranking/{type}",                      "GET",  "랭킹"],
        ["/api/trader/positions",                    "GET",  "보유 포지션"],
        ["/api/trader/orders",                       "GET",  "주문 이력"],
        ["/api/themes/*",                            "GET",  "테마 정보"],
        ["/api/disclosures",                         "GET",  "공시 목록"],
        ["/api/news",                                "GET",  "뉴스 목록"],
        ["/api/intel/summary",                       "GET",  "인텔리전스 요약"],
        ["/api/settings/*",                          "*",    "설정 관리"],
    ],
    col_widths=[6.5, 2.0, 8.0],
)

h(2, "부록 B. 주요 환경변수")
table(
    ["환경변수", "기본값", "설명"],
    [
        ["POSTGRES_DB / _USER / _PASSWORD", "-",       "DB 접속 정보"],
        ["POSTGRES_PORT",                   "5432",    "DB 포트"],
        ["REDIS_PORT",                      "6379",    "Redis 포트"],
        ["API_PORT",                        "8000",    "API + 프론트엔드 포트"],
        ["ML_API_PORT",                     "8001",    "ML 서비스 포트"],
        ["NOTIFIER_PORT",                   "8003",    "알림 서비스 포트"],
        ["TRADER_PORT",                     "8004",    "자동매매 서비스 포트"],
        ["KIS_APP_KEY / _APP_SECRET",       "-",       "KIS OpenAPI 인증"],
        ["KIS_ACCOUNT_NO / _CD",            "-",       "계좌번호/상품코드"],
        ["KIS_MODE",                        "paper",   "paper(모의) / live(실전)"],
        ["DART_API_KEY",                    "-",       "금감원 DART API 키"],
        ["KIND_API_KEY",                    "-",       "KIND API 키"],
        ["GOVDATA_API_KEY",                 "-",       "금융위 시가총액 API 키"],
        ["TELEGRAM_BOT_TOKEN / CHAT_ID",    "-",       "텔레그램 알림"],
        ["REC_COOLDOWN_MINUTES",            "60",      "동일 종목 재추천 쿨다운(분)"],
        ["REC_PERF_MIN_PROB",               "0.236",   "성과 추적 최소 확률"],
        ["REC_RECOVERY_HOURS",              "24",      "실패 후 회복 대기(시간)"],
        ["REGIME_BEAR_THRESHOLD",           "-3.0",    "약세 국면 임계값(%)"],
        ["REGIME_BULL_THRESHOLD",           "-0.5",    "강세 국면 임계값(%)"],
        ["REGIME_BEAR_SUPPRESS",            "true",    "약세 국면 추천 억제"],
        ["REGIME_BEAR_PROB_MULT",           "0.75",    "약세 국면 확률 배수"],
        ["REGIME_NEUTRAL_PROB_MULT",        "0.88",    "중립 국면 확률 배수"],
        ["ML_MIN_AUC_THRESHOLD",            "0.57",    "재학습 배포 최소 AUC"],
        ["ML_RETRAIN_DAYS",                 "14,28",   "자동 재학습 일자(월중)"],
        ["ML_RETRAIN_HOUR",                 "0",       "자동 재학습 시각(KST)"],
        ["ML_RETRAIN_TRAIN_YEARS",          "8",       "학습 기간(년)"],
        ["ML_RETRAIN_MAX_CODES",            "500",     "재학습 대상 종목 수"],
        ["BARS_BACKFILL_DAYS",              "3285",    "일봉 백필 일수(9년)"],
    ],
    col_widths=[5.5, 3.0, 8.0],
)

h(2, "부록 C. 지원 이벤트 타입 (14종)")
p("EVENT_TYPES = [")
p("  'VOLUME_SURGE', 'AMOUNT_SURGE',")
p("  'BREAKOUT_52W', 'BREAKOUT_26W', 'BREAKOUT_13W', 'BREAKOUT_20D',")
p("  'VI_TRIGGERED', 'LONG_WHITE_CANDLE', 'HAMMER_CANDLE', 'MORNING_STAR',")
p("  'SUPPLY_ANOMALY', 'POST_DISCLOSURE_SURGE',")
p("  'SHORT_SURGE', 'DUAL_BUY_STREAK',")
p("]")

h(2, "부록 D. Redis Pub/Sub 채널")
table(
    ["채널", "발행자", "구독자", "페이로드"],
    [
        ["ch:feature",        "detector",    "analyzer, api", "feature_event_id, code, event_type, signal_score"],
        ["ch:recommendation", "recommender", "trader, notifier, api", "rec_id, code, action, success_prob, prices"],
        ["ch:trade_result",   "trader",      "recommender, api", "order_id, code, status, filled_qty, price"],
        ["ch:alert",          "notifier",    "-",             "channel, target, message, ts"],
    ],
    col_widths=[3.5, 3.0, 4.5, 5.0],
)

h(2, "부록 E. Docker Compose 프로파일")
table(
    ["프로파일", "실행 명령", "용도"],
    [
        ["(기본)",     "docker compose up -d",                          "상시 운영 서비스"],
        ["backfill",   "docker compose --profile backfill up",          "9년치 일봉 재백필"],
        ["tools",      "docker compose --profile tools run --rm ml-train", "ML 학습·수동 백필"],
        ["monitoring", "docker compose --profile monitoring up -d",     "Prometheus + Grafana"],
    ],
    col_widths=[3.0, 8.0, 5.0],
)

# ═══════════════════════════════════════════════════════════
# 저장
# ═══════════════════════════════════════════════════════════
doc.save(OUT)
size_kb = os.path.getsize(OUT) / 1024
print(f"OK: 생성 완료 → {OUT}")
print(f"크기: {size_kb:.1f} KB")
