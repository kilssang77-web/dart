import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

d = Document("D:/a2m/atom-harness-base-Dart/kospi-feature-stock/Quant_Eye_사용자매뉴얼_v1.0.docx")

print(f"총 단락 수: {len(d.paragraphs)}")
print(f"총 표 수: {len(d.tables)}")

# 그림 개수
img_count = 0
for shape in d.inline_shapes:
    img_count += 1
print(f"인라인 그림 수: {img_count}")

# 헤딩 목록
print("\n=== 헤딩 구조 ===")
for p in d.paragraphs:
    if p.style and p.style.name.startswith("Heading"):
        level = p.style.name.replace("Heading ", "")
        indent = "  " * (int(level) - 1)
        print(f"{indent}[H{level}] {p.text[:80]}")
